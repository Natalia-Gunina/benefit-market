#!/usr/bin/env python3
"""Generate technical specification DOCX for ВКР."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup (A4, ГОСТ margins) ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# ── Style helpers ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(14)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(0)
style_normal.paragraph_format.space_before = Pt(0)

def set_font(run, size=14, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_text(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    else:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=14 if level > 1 else 16, bold=True)
    return p

def add_para(text, bold=False, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, bold=bold)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=12)
    doc.add_paragraph()  # spacing after table
    return table


# ═══════════════════════════════════════════════════════════════
# CHAPTER 3
# ═══════════════════════════════════════════════════════════════

add_heading_text('Глава 3. Проектирование и разработка информационной системы «Benefit Market»')

# ── 3.1 ──
add_heading_text('3.1. Архитектура приложения')

# 3.1.1
add_heading_text('3.1.1. Текущая архитектура', level=2)

add_para(
    'Система реализована как монолитное full-stack веб-приложение на базе '
    'фреймворка Next.js 16 с использованием паттерна App Router. Выбор монолитной '
    'архитектуры обусловлен масштабом MVP-версии продукта: единая кодовая база '
    'упрощает разработку, отладку и развёртывание на начальном этапе.'
)

add_para('Приложение разделено на три логических слоя:')

add_para(
    '1) Презентационный слой — React-компоненты с разделением на серверные '
    '(загрузка данных, SSR) и клиентские (интерактивность). Управление серверным '
    'состоянием осуществляется через TanStack React Query, клиентским — через Zustand.'
)
add_para(
    '2) Слой бизнес-логики — сервисы (order.service, analytics.service, '
    'accrual.service), доменные правила (evaluator элигибильности), '
    'Zod-валидация входных данных на границах системы.'
)
add_para(
    '3) Слой данных — PostgreSQL с Row-Level Security (RLS), обеспечивающей '
    'мультитенантную изоляцию на уровне СУБД. Аутентификация через JWT-токены '
    'Supabase Auth. 56 API-маршрутов реализованы как Next.js Route Handlers.'
)

add_para(
    'Взаимодействие между слоями: клиент отправляет HTTPS/JSON-запросы к серверному '
    'слою Next.js, который через Middleware выполняет аутентификацию и проверку ролей '
    '(RBAC). Далее запрос передаётся в сервисный слой, обращающийся к PostgreSQL через '
    'Supabase-клиент. В продуктивной среде перед Next.js стоит Caddy в роли '
    'reverse-proxy с автоматическим SSL-терминированием.'
)

# 3.1.2
add_heading_text('3.1.2. Целевая архитектура', level=2)

add_para(
    'На этапе масштабирования планируется переход к модульному монолиту '
    'с выделением bounded contexts. Переход выполняется без изменения инфраструктуры: '
    'модули выделяются как пакеты внутри монорепозитория, взаимодействуют через '
    'типизированные интерфейсы. При необходимости горизонтального масштабирования '
    'модули могут быть вынесены в отдельные микросервисы.'
)

add_table(
    ['Модуль', 'Ответственность', 'Критерий выделения'],
    [
        ['benefits-core', 'Каталог льгот, правила элигибильности', 'Высокая частота изменений бизнес-правил'],
        ['marketplace', 'Провайдеры, офферинги, отзывы', 'Независимый цикл разработки'],
        ['billing', 'Кошельки, начисления, списания', 'Транзакционная целостность'],
        ['identity', 'Аутентификация, RBAC, профили', 'Интеграция с корпоративным SSO'],
        ['integration', 'Коннекторы HRIS / 1С', 'Изоляция внешних зависимостей'],
    ]
)

# 3.1.3
add_heading_text('3.1.3. Стек технологий и обоснование выбора', level=2)

add_para('Серверная часть:', bold=True, indent=False)
add_table(
    ['Технология', 'Версия', 'Обоснование'],
    [
        ['Next.js (App Router)', '16.1.6', 'Гибридный рендеринг (SSR/CSR), встроенные API Routes, Turbopack'],
        ['TypeScript', '5.x (strict)', 'Сквозная типизация от БД до UI, Zod-валидация на границах'],
        ['Supabase', '2.97.0', 'Управляемый PostgreSQL, JWT-аутентификация, RLS-политики'],
        ['PostgreSQL', '15+', 'JSONB, Row-Level Security, расширяемость'],
        ['Node.js', '20 LTS', 'Долгосрочная поддержка, единый язык клиент/сервер'],
    ]
)

add_para('Клиентская часть:', bold=True, indent=False)
add_table(
    ['Технология', 'Версия', 'Обоснование'],
    [
        ['React', '19.2.3', 'Server Components, Concurrent-рендеринг'],
        ['shadcn/ui + Radix UI', '1.4.3', 'Доступность (WCAG 2.1), headless-компоненты'],
        ['Tailwind CSS', '4.x', 'Utility-first, tree-shaking, минимальный CSS-бандл'],
        ['TanStack React Query', '5.90.21', 'Кеширование, фоновая ревалидация, инвалидация'],
        ['Zustand', '5.0.11', 'Минимальный стор с persist-middleware для корзины'],
        ['Recharts', '3.7.0', 'Декларативные React-графики для аналитики'],
    ]
)

add_para('Инфраструктура:', bold=True, indent=False)
add_table(
    ['Технология', 'Обоснование'],
    [
        ['Docker (multi-stage, Alpine)', 'Воспроизводимые сборки, standalone-образ ~50 МБ'],
        ['Caddy', 'Автоматический HTTPS (Let\'s Encrypt), HTTP/2'],
        ['GitHub Actions', 'CI/CD: lint → typecheck → test → build → deploy'],
        ['Vitest', 'Быстрое выполнение тестов, совместимость с TypeScript'],
    ]
)

# 3.1.4
add_heading_text('3.1.4. Схема базы данных', level=2)

add_para(
    'База данных содержит 18 таблиц, 8 пользовательских ENUM-типов и '
    '37 RLS-политик. Данные организованы в 5 логических групп.'
)

add_table(
    ['Группа', 'Таблицы', 'Назначение'],
    [
        ['Ядро', 'tenants, users, employee_profiles', 'Мультитенантность, учётные записи, атрибуты сотрудников'],
        ['Льготы', 'benefits, benefit_categories, eligibility_rules', 'Каталог с JSONB-правилами элигибильности'],
        ['Финансы', 'budget_policies, wallets, point_ledger', 'Начисление баллов, резервирование, аудит операций'],
        ['Заказы', 'orders, order_items', 'Жизненный цикл: pending → reserved → paid / cancelled / expired'],
        ['Маркетплейс', 'providers, provider_offerings, tenant_offerings, reviews, provider_users', 'Провайдеры, модерация, подключение к тенантам'],
        ['Аудит', 'audit_log', 'Журнал действий с JSONB-diff'],
    ]
)

add_para('ENUM-типы, используемые в базе данных:', bold=True, indent=False)
add_table(
    ['Тип', 'Значения'],
    [
        ['user_role', 'employee, hr, admin, provider'],
        ['order_status', 'pending, reserved, paid, cancelled, expired'],
        ['ledger_type', 'accrual, reserve, spend, release, expire'],
        ['budget_period', 'monthly, quarterly, yearly'],
        ['provider_status', 'pending, verified, suspended, rejected'],
        ['offering_status', 'draft, pending_review, published, archived'],
        ['review_status', 'visible, hidden, flagged'],
        ['provider_user_role', 'owner, admin, member'],
    ]
)

# 3.1.5
add_heading_text('3.1.5. Интеграции (HRIS / 1С)', level=2)

add_para(
    'На текущем этапе реализован импорт сотрудников из CSV-файлов '
    '(/api/import/employees) с парсингом через библиотеку PapaParser. '
    'Это покрывает сценарий первичной загрузки данных из внешних систем.'
)

add_para('Планируемые интеграции:', bold=True, indent=False)
add_table(
    ['Система', 'Способ интеграции', 'Данные', 'Приоритет'],
    [
        ['1С:ЗУП', 'REST API / выгрузка CSV', 'Сотрудники, грейды, подразделения', 'Высокий'],
        ['HRIS (SAP SF, BambooHR)', 'Webhook / REST API', 'Профили, орг. структура', 'Средний'],
        ['SSO (SAML 2.0 / OIDC)', 'Supabase Auth Provider', 'Аутентификация через корп. IdP', 'Высокий'],
        ['Платёжные системы', 'REST API', 'Подтверждение оплаты, статусы', 'Средний'],
    ]
)

add_para(
    'Архитектурно интеграции будут изолированы в модуле integration с паттерном '
    'Adapter, что позволит подключать новые внешние системы без изменения ядра.'
)

# ── 3.2 ──
add_heading_text('3.2. Оценка и планирование разработки')

# 3.2.1
add_heading_text('3.2.1. Методология разработки', level=2)

add_para(
    'Разработка ведётся по методологии Agile с использованием фреймворка Scrum. '
    'Длительность спринта — 2 недели. Управление задачами осуществляется через '
    'GitHub Issues и Beads (git-backed issue tracker). Контроль версий — Git с '
    'trunk-based development и feature-ветками.'
)

add_table(
    ['Параметр', 'Значение'],
    [
        ['Длительность спринта', '2 недели'],
        ['Инструмент управления', 'GitHub Issues + Beads'],
        ['Контроль версий', 'Git (GitHub), trunk-based development'],
        ['CI/CD', 'GitHub Actions (lint → typecheck → test → build → deploy)'],
        ['Ревью кода', 'Pull Request с обязательным прохождением CI'],
    ]
)

add_para('Scrum-церемонии:', bold=True, indent=False)
add_para('— Sprint Planning: начало спринта, декомпозиция задач из бэклога;')
add_para('— Daily Standup: синхронизация по блокерам;')
add_para('— Sprint Review: демонстрация результатов спринта;')
add_para('— Retrospective: анализ процесса и улучшения.')

# 3.2.2
add_heading_text('3.2.2. Состав команды и роли', level=2)

add_table(
    ['Роль', 'Кол-во', 'Ответственность'],
    [
        ['Product Owner', '1', 'Приоритизация бэклога, приёмка результатов'],
        ['Full-stack разработчик', '1', 'Реализация фронтенда и бэкенда, код-ревью'],
        ['UI/UX дизайнер', '1', 'Макеты в Figma, дизайн-система'],
        ['Научный руководитель', '1', 'Контроль соответствия требованиям ВКР'],
    ]
)

# 3.2.3
add_heading_text('3.2.3. Декомпозиция задач: бэклог по эпикам', level=2)

add_para(
    'Бэклог проекта декомпозирован на 6 эпиков, каждый из которых покрывает '
    'функциональную область системы. Оценка трудозатрат выполнена в Story Points '
    'по шкале Фибоначчи.'
)

add_table(
    ['#', 'Эпик', 'Story Points', 'Спринт', 'Статус'],
    [
        ['E1', 'Инфраструктура и каркас', '13', '1', 'Завершён'],
        ['', '— Инициализация проекта (Next.js, TS, Tailwind)', '3', '', ''],
        ['', '— Схема БД и миграции (18 таблиц, RLS)', '5', '', ''],
        ['', '— Аутентификация (Supabase Auth, RBAC)', '5', '', ''],
        ['E2', 'Ядро системы льгот', '21', '2–3', 'Завершён'],
        ['', '— Каталог льгот (CRUD, категории, фильтрация)', '5', '', ''],
        ['', '— Правила элигибильности (JSONB, evaluator)', '8', '', ''],
        ['', '— Бюджетные политики и начисление баллов', '5', '', ''],
        ['', '— Кошельки и журнал операций', '3', '', ''],
        ['E3', 'Заказы и корзина', '13', '3–4', 'Завершён'],
        ['', '— Корзина (Zustand, localStorage persist)', '3', '', ''],
        ['', '— Жизненный цикл заказа', '5', '', ''],
        ['', '— Резервирование и списание баллов', '3', '', ''],
        ['', '— Автоэкспирация заказов (15 мин)', '2', '', ''],
        ['E4', 'Маркетплейс провайдеров', '21', '4–5', 'Завершён'],
        ['', '— Регистрация и модерация провайдеров', '5', '', ''],
        ['', '— CRUD офферингов (draft → published)', '5', '', ''],
        ['', '— Подключение офферингов к тенанту', '3', '', ''],
        ['', '— Система отзывов', '3', '', ''],
        ['', '— Панель провайдера (аналитика, команда)', '5', '', ''],
        ['E5', 'Дашборды и аналитика', '13', '5–6', 'Завершён'],
        ['', '— Дашборд сотрудника', '3', '', ''],
        ['', '— Дашборд HR (Recharts)', '5', '', ''],
        ['', '— Дашборд администратора', '3', '', ''],
        ['', '— Импорт сотрудников из CSV', '2', '', ''],
        ['E6', 'Инфраструктура деплоя', '8', '6', 'Завершён'],
        ['', '— Docker multi-stage build', '3', '', ''],
        ['', '— CI/CD (GitHub Actions)', '3', '', ''],
        ['', '— Продуктивная среда (Caddy, VPS)', '2', '', ''],
        ['', 'Итого', '89 SP', '6 спринтов', ''],
    ]
)

# 3.2.4
add_heading_text('3.2.4. Roadmap', level=2)

add_table(
    ['Спринт', 'Период', 'Содержание'],
    [
        ['1', 'Недели 1–2', 'E1: Инфраструктура и каркас'],
        ['2', 'Недели 3–4', 'E2: Ядро системы льгот (ч.1)'],
        ['3', 'Недели 5–6', 'E2 + E3: Ядро (ч.2) + Заказы'],
        ['4', 'Недели 7–8', 'E3 + E4: Заказы + Маркетплейс (ч.1)'],
        ['5', 'Недели 9–10', 'E4 + E5: Маркетплейс (ч.2) + Дашборды'],
        ['6', 'Недели 11–12', 'E5 + E6: Дашборды + Деплой'],
    ]
)

add_para('Следующий этап развития (после MVP):', bold=True, indent=False)
add_para('— интеграция с 1С:ЗУП и корпоративными HRIS-системами;')
add_para('— SSO через SAML 2.0 / OpenID Connect;')
add_para('— интернационализация (i18n) интерфейса;')
add_para('— переход к модульному монолиту;')
add_para('— мобильное приложение (React Native);')
add_para('— нагрузочное тестирование и оптимизация.')

# 3.2.5
add_heading_text('3.2.5. Метрики проекта', level=2)

add_table(
    ['Метрика', 'Значение'],
    [
        ['API-маршрутов', '56'],
        ['Таблиц в БД', '18'],
        ['RLS-политик безопасности', '37'],
        ['UI-компонентов', '26'],
        ['React-хуков', '7'],
        ['Сервисов бизнес-логики', '4'],
        ['Миграций БД', '5'],
        ['GitHub Actions workflows', '2'],
        ['Production-зависимостей', '20'],
        ['Ролей пользователей', '4'],
    ]
)

# ── Save ──
output_path = 'docs/technical-specification.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
